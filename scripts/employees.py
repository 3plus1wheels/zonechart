import pandas as pd
import numpy as np
from docx import Document

class employee:
    def __init__(self, name, position, mens, womens, cash, fits, stock, greet, replen, sfs):
        self.name = name
        self.position = position
        self.mens = mens
        self.womens = womens
        self.cash = cash
        self.fits = fits
        self.stock = stock
        self.greet = greet
        self.replen = replen
        self.sfs = sfs
        self.info = self.get_info()

    doc = Document("data/EMPLOYEE  KNOWLEDGE BASE NEW.docx")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                print(cell.text, end="\t")
            print()

    table = doc.tables[0]

    


    print(table.cell(1, 0).text)  # Example to access the first cell of the first table
    

    def get_info(self):
        return f"Employee Name: {self.name}, Position: {self.position}, Mens: {self.mens}, Womens: {self.womens}, Cash: {self.cash}, Fits: {self.fits}, Stock: {self.stock}, Greet: {self.greet}"
    
    